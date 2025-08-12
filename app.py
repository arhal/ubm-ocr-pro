import os
import base64
import tempfile
import io
import re
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
from mistralai import Mistral
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Configure upload settings
UPLOAD_FOLDER = '/tmp' if os.environ.get('VERCEL') else 'uploads'
ALLOWED_EXTENSIONS = {'pdf'}

# Create uploads directory only if not on Vercel
if not os.environ.get('VERCEL') and not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_pdf_ocr(pdf_path):
    """Process PDF using Mistral OCR from file path"""
    try:
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        return process_pdf_ocr_from_bytes(pdf_bytes)
    except Exception as e:
        raise Exception(f"OCR processing failed: {str(e)}")

def process_pdf_ocr_from_bytes(pdf_bytes):
    """Process PDF using Mistral OCR from bytes"""
    try:
        b64 = base64.b64encode(pdf_bytes).decode()
        api_key = os.environ.get('MISTRAL_API_KEY', '97ZQlsV45YrDusgZRwjArWGbh3nerFPb')
        client = Mistral(api_key=api_key)
        
        resp = client.ocr.process(
            model="mistral-ocr-latest",
            document={
                "type": "document_url",
                "document_url": f"data:application/pdf;base64,{b64}"
            },
            include_image_base64=True,
        )
        
        # Combine all pages
        full_text = ""
        for page in resp.pages:
            full_text += page.markdown + "\n\n"
        
        return full_text.strip()
    
    except Exception as e:
        raise Exception(f"OCR processing failed: {str(e)}")

def create_word_document(text, filename):
    """Create a Word document from markdown text"""
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading('OCR Results', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add source info
        info_para = doc.add_paragraph()
        info_para.add_run('Source File: ').bold = True
        info_para.add_run(filename)
        info_para.add_run('\nProcessed by: ').bold = True
        info_para.add_run('UBM OCR Pro')
        
        # Add separator
        doc.add_paragraph('_' * 50)
        
        # Process the markdown text
        lines = text.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - add paragraph break
                if current_paragraph:
                    current_paragraph = None
                continue
            
            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                if header_text:
                    doc.add_heading(header_text, level)
                current_paragraph = None
                continue
            
            # Handle bold text **text**
            if '**' in line:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        current_paragraph.add_run(part)
                    else:
                        current_paragraph.add_run(part).bold = True
                current_paragraph.add_run('\n')
                continue
            
            # Regular text
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            
            current_paragraph.add_run(line + '\n')
        
        # Save to memory
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
    
    except Exception as e:
        raise Exception(f"Failed to create Word document: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        try:
            # For Vercel, process file directly from memory
            if os.environ.get('VERCEL'):
                # Process PDF directly from uploaded file object
                pdf_bytes = file.read()
                result = process_pdf_ocr_from_bytes(pdf_bytes)
            else:
                # Save uploaded file temporarily for local/other deployments
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                
                # Process OCR
                result = process_pdf_ocr(filepath)
                
                # Clean up uploaded file
                os.remove(filepath)
            
            return jsonify({
                'success': True,
                'result': result,
                'filename': file.filename
            })
            
        except Exception as e:
            # Clean up file if it exists (for non-Vercel deployments)
            if not os.environ.get('VERCEL') and 'filepath' in locals() and os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'error': str(e)}), 500
    
    return jsonify({'error': 'Invalid file type. Please upload a PDF file.'}), 400

@app.route('/download-docx', methods=['POST'])
def download_docx():
    try:
        data = request.get_json()
        text = data.get('text', '')
        filename = data.get('filename', 'ocr_result')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Create Word document
        doc_io = create_word_document(text, filename)
        
        # Generate filename
        docx_filename = filename.replace('.pdf', '_ocr_result.docx')
        if not docx_filename.endswith('.docx'):
            docx_filename += '_ocr_result.docx'
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=docx_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    debug = os.environ.get('FLASK_ENV') != 'production'
    app.run(debug=debug, host='0.0.0.0', port=port)